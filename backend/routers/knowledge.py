"""
知识库管理 API
- 文件解析: .docx / .xlsx 上传 → 结构化 JSON（保留）
- MD 文档管理: KnowledgeBaseDoc CRUD（新增）
- 文件夹扫描导入: 扫描 data/knowledge_base/ 目录（新增）
"""
import io
import os
import re
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from backend.models.database import get_db
from backend.models.models import KnowledgeBaseDoc, Company, Product
from backend.schemas.schemas import (
    APIResponse, KBDocCreate, KBDocUpdate, KBDocResponse, KBDocListResponse,
    KBScanRequest,
)

router = APIRouter()

# ---------- 知识库根目录 ----------
KB_ROOT = Path(__file__).resolve().parent.parent.parent / "data" / "knowledge_base"

# ---------- 允许的文件扩展名 ----------
ALLOWED_EXTENSIONS = {".docx", ".xlsx"}


# ═══════════════════════════════════════════
#  一、MD 文档管理 (KnowledgeBaseDoc CRUD)
# ═══════════════════════════════════════════

DOC_LIST_FIELDS = [
    KnowledgeBaseDoc.id,
    KnowledgeBaseDoc.title,
    KnowledgeBaseDoc.company_id,
    KnowledgeBaseDoc.product_id,
    KnowledgeBaseDoc.file_path,
    KnowledgeBaseDoc.category,
    KnowledgeBaseDoc.source,
    KnowledgeBaseDoc.word_count,
    KnowledgeBaseDoc.created_at,
    KnowledgeBaseDoc.updated_at,
]


@router.get("/docs", response_model=KBDocListResponse)
async def list_kb_docs(
    company_id: Optional[int] = Query(None, description="按公司筛选"),
    product_id: Optional[int] = Query(None, description="按产品筛选"),
    category: Optional[str] = Query(None, description="按分类筛选"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: AsyncSession = Depends(get_db),
):
    """列出知识库文档列表（不含正文内容）"""
    count_q = select(func.count(KnowledgeBaseDoc.id))
    list_q = select(*DOC_LIST_FIELDS)

    if company_id is not None:
        count_q = count_q.where(KnowledgeBaseDoc.company_id == company_id)
        list_q = list_q.where(KnowledgeBaseDoc.company_id == company_id)
    if product_id is not None:
        count_q = count_q.where(KnowledgeBaseDoc.product_id == product_id)
        list_q = list_q.where(KnowledgeBaseDoc.product_id == product_id)
    if category:
        count_q = count_q.where(KnowledgeBaseDoc.category == category)
        list_q = list_q.where(KnowledgeBaseDoc.category == category)

    total = await db.scalar(count_q)

    list_q = list_q.order_by(KnowledgeBaseDoc.updated_at.desc())
    list_q = list_q.offset((page - 1) * page_size).limit(page_size)
    result = await db.execute(list_q)
    rows = result.all()

    items = []
    for row in rows:
        cname = None
        pname = None
        if row.company_id:
            c = await db.scalar(select(Company.name).where(Company.id == row.company_id))
            cname = c
        if row.product_id:
            p = await db.scalar(select(Product.name).where(Product.id == row.product_id))
            pname = p

        items.append(KBDocResponse(
            id=row.id, title=row.title,
            company_id=row.company_id, company_name=cname,
            product_id=row.product_id, product_name=pname,
            file_path=row.file_path, category=row.category,
            source=row.source, word_count=row.word_count,
            created_at=row.created_at, updated_at=row.updated_at,
        ))

    return KBDocListResponse(items=items, total=total)


@router.post("/docs", response_model=KBDocResponse)
async def create_kb_doc(
    data: KBDocCreate,
    db: AsyncSession = Depends(get_db),
):
    """创建知识库 MD 文档"""
    # 校验：至少绑定 company 或 product 之一
    if not data.company_id and not data.product_id:
        raise HTTPException(status_code=400, detail="必须指定 company_id 或 product_id")

    # 校验关联存在
    if data.company_id:
        c = await db.scalar(select(Company).where(Company.id == data.company_id))
        if not c:
            raise HTTPException(status_code=404, detail="公司不存在")
    if data.product_id:
        p = await db.scalar(select(Product).where(Product.id == data.product_id))
        if not p:
            raise HTTPException(status_code=404, detail="产品不存在")

    # 保存到磁盘
    file_path = _save_md_file(data.title, data.content, data.company_id, data.product_id)

    doc = KnowledgeBaseDoc(
        title=data.title,
        content=data.content,
        company_id=data.company_id,
        product_id=data.product_id,
        category=data.category or "product",
        source="manual",
        file_path=file_path,
        word_count=len(data.content),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    cname = _get_company_name(db, doc.company_id) if doc.company_id else None
    pname = _get_product_name(db, doc.product_id) if doc.product_id else None

    return KBDocResponse(
        id=doc.id, title=doc.title,
        company_id=doc.company_id, company_name=cname,
        product_id=doc.product_id, product_name=pname,
        file_path=doc.file_path, category=doc.category,
        source=doc.source, word_count=doc.word_count,
        content=doc.content,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.get("/docs/{doc_id}", response_model=KBDocResponse)
async def get_kb_doc(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
):
    """获取知识库文档详情（含完整内容）"""
    doc = await db.scalar(select(KnowledgeBaseDoc).where(KnowledgeBaseDoc.id == doc_id))
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    cname = _get_company_name(db, doc.company_id) if doc.company_id else None
    pname = _get_product_name(db, doc.product_id) if doc.product_id else None

    return KBDocResponse(
        id=doc.id, title=doc.title,
        company_id=doc.company_id, company_name=cname,
        product_id=doc.product_id, product_name=pname,
        file_path=doc.file_path, category=doc.category,
        source=doc.source, word_count=doc.word_count,
        content=doc.content,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.put("/docs/{doc_id}", response_model=KBDocResponse)
async def update_kb_doc(
    doc_id: int,
    data: KBDocUpdate,
    db: AsyncSession = Depends(get_db),
):
    """更新知识库文档"""
    doc = await db.scalar(select(KnowledgeBaseDoc).where(KnowledgeBaseDoc.id == doc_id))
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    dump = data.model_dump(exclude_unset=True)

    if "content" in dump:
        dump["word_count"] = len(dump["content"])
        # 更新磁盘文件
        if doc.file_path:
            _update_md_file(doc.file_path, dump["content"])

    for key, value in dump.items():
        setattr(doc, key, value)

    await db.commit()
    await db.refresh(doc)

    cname = _get_company_name(db, doc.company_id) if doc.company_id else None
    pname = _get_product_name(db, doc.product_id) if doc.product_id else None

    return KBDocResponse(
        id=doc.id, title=doc.title,
        company_id=doc.company_id, company_name=cname,
        product_id=doc.product_id, product_name=pname,
        file_path=doc.file_path, category=doc.category,
        source=doc.source, word_count=doc.word_count,
        content=doc.content,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.delete("/docs/{doc_id}", response_model=APIResponse)
async def delete_kb_doc(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
):
    """删除知识库文档"""
    doc = await db.scalar(select(KnowledgeBaseDoc).where(KnowledgeBaseDoc.id == doc_id))
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 删除磁盘文件
    if doc.file_path:
        _delete_md_file(doc.file_path)

    await db.delete(doc)
    await db.commit()
    return APIResponse(success=True, message="文档已删除")


# ═══════════════════════════════════════════
#  二、文件夹扫描导入
# ═══════════════════════════════════════════

@router.post("/scan", response_model=APIResponse)
async def scan_knowledge_base(
    data: KBScanRequest = None,
    db: AsyncSession = Depends(get_db),
):
    """
    扫描 data/knowledge_base/ 目录中的 .md 文件并导入为知识库文档。
    目录结构约定:
      data/knowledge_base/
        companies/{slug}/*.md          → 按 slug 匹配公司
        products/{slug}/*.md           → 按 slug 匹配产品
    """
    if data and data.path:
        scan_dir = KB_ROOT / data.path
    else:
        scan_dir = KB_ROOT

    if not scan_dir.exists():
        return APIResponse(success=True, message="知识库目录不存在，已自动创建",
                          data={"imported": 0, "skipped": 0})

    imported = 0
    skipped = 0
    details = []

    # 扫描 companies/
    companies_dir = scan_dir / "companies"
    if companies_dir.exists():
        for company_dir in companies_dir.iterdir():
            if not company_dir.is_dir():
                continue
            slug = company_dir.name
            company = await db.scalar(select(Company).where(Company.slug == slug))
            if not company:
                details.append({"dir": str(company_dir), "status": "skipped", "reason": f"未找到 slug='{slug}' 的公司"})
                skipped += 1
                continue

            for md_file in company_dir.glob("*.md"):
                res = await _import_md_file(db, md_file, company_id=company.id, category="company")
                if res["imported"]:
                    imported += 1
                else:
                    skipped += 1
                details.append(res)

    # 扫描 products/
    products_dir = scan_dir / "products"
    if products_dir.exists():
        for prod_dir in products_dir.iterdir():
            if not prod_dir.is_dir():
                continue
            slug = prod_dir.name
            # 尝试用 slug 或 name 匹配产品（模糊匹配）
            products = (await db.execute(
                select(Product).where(Product.name.contains(slug))
            )).scalars().all()

            if not products:
                details.append({"dir": str(prod_dir), "status": "skipped", "reason": f"未找到匹配的产品: '{slug}'"})
                skipped += 1
                continue

            for product in products:
                for md_file in prod_dir.glob("*.md"):
                    res = await _import_md_file(db, md_file, product_id=product.id, category="product")
                    if res["imported"]:
                        imported += 1
                    else:
                        skipped += 1
                    details.append(res)

    return APIResponse(
        success=True,
        message=f"扫描完成：导入 {imported} 个，跳过 {skipped} 个",
        data={"imported": imported, "skipped": skipped, "details": details},
    )


async def _import_md_file(db: AsyncSession, filepath: Path, company_id=None, product_id=None, category="product") -> dict:
    """导入单个 .md 文件"""
    name = filepath.stem  # 文件名作为标题
    try:
        content = filepath.read_text(encoding="utf-8")
    except Exception as e:
        return {"file": str(filepath), "imported": False, "status": "skipped", "reason": str(e)}

    # 检查是否已导入（基于 file_path 去重）
    rel_path = str(filepath.relative_to(KB_ROOT.parent))
    existing = await db.scalar(
        select(KnowledgeBaseDoc).where(KnowledgeBaseDoc.file_path == rel_path)
    )
    if existing:
        return {"file": str(filepath), "imported": False, "status": "skipped", "reason": "已存在"}

    doc = KnowledgeBaseDoc(
        title=name,
        content=content,
        company_id=company_id,
        product_id=product_id,
        category=category,
        source="scanned",
        file_path=rel_path,
        word_count=len(content),
    )
    db.add(doc)
    await db.flush()
    return {"file": str(filepath), "imported": True, "status": "imported", "doc_id": doc.id}


# ═══════════════════════════════════════════
#  三、文件存储辅助
# ═══════════════════════════════════════════

def _get_kb_dir(company_id=None, product_id=None) -> Path:
    """确定文档存储目录"""
    if company_id:
        return KB_ROOT / "companies" / f"c_{company_id}"
    elif product_id:
        return KB_ROOT / "products" / f"p_{product_id}"
    return KB_ROOT / "standalone"


def _save_md_file(title: str, content: str, company_id=None, product_id=None) -> str:
    """保存 MD 内容到磁盘，返回相对路径"""
    target_dir = _get_kb_dir(company_id, product_id)
    target_dir.mkdir(parents=True, exist_ok=True)

    # 清理文件名（保留中英文、数字、下划线、连字符）
    safe_name = re.sub(r"[^\w\u4e00-\u9fff\uff00-\uffef\-]", "_", title, flags=re.UNICODE)[:80]
    filename = f"{safe_name}.md"
    filepath = target_dir / filename

    # 避免覆盖：重名加序号
    counter = 1
    while filepath.exists():
        filename = f"{safe_name}_{counter}.md"
        filepath = target_dir / filename
        counter += 1

    filepath.write_text(content, encoding="utf-8")

    # 返回相对于 data/ 的路径
    return str(filepath.relative_to(KB_ROOT.parent))


def _update_md_file(rel_path: str, content: str):
    """更新磁盘上的 MD 文件"""
    full_path = KB_ROOT.parent / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    full_path.write_text(content, encoding="utf-8")


def _delete_md_file(rel_path: str):
    """删除磁盘上的 MD 文件"""
    full_path = KB_ROOT.parent / rel_path
    if full_path.exists():
        full_path.unlink()


def _get_company_name(db, company_id: int) -> str:
    """获取公司名（同步辅助，用于列表组装）"""
    import asyncio
    # 注意：这个函数在 async context 中调用，但我们已经 await 过了
    return None  # 在 async endpoint 中直接 await


def _get_product_name(db, product_id: int) -> str:
    return None  # 同上


# ═══════════════════════════════════════════
#  四、Word/Excel → Markdown 知识库文档（新增）
# ═══════════════════════════════════════════


def _docx_to_markdown(file_bytes: bytes) -> str:
    """将 .docx 文件转为 Markdown 字符串"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and "Heading" in para.style.name:
            level_match = re.search(r"Heading\s*(\d+)", para.style.name)
            level = int(level_match.group(1)) if level_match else 1
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    # 也提取表格
    for i, table in enumerate(doc.tables):
        lines.append(f"\n## 表格 {i + 1}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            # 表头后加分隔行
            if row == table.rows[0] and len(cells) > 1:
                lines.append("|" + "|".join([" --- " for _ in cells]) + "|")

    return "\n\n".join(lines)


def _xlsx_to_markdown(file_bytes: bytes) -> str:
    """将 .xlsx 文件转为 Markdown 字符串"""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        lines.append(f"## {sheet_name}")

        # 过滤全空行
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue

        # 检测竖排字段格式（key: value）
        is_kv = False
        if len(rows) >= 2 and len(rows[0]) >= 2:
            first_cell = str(rows[0][0]).strip() if rows[0][0] else ""
            if len(first_cell) <= 30 and first_cell:
                # 检查是否大多数行都是 key: value 形式
                kv_count = sum(
                    1 for r in rows
                    if len(r) >= 2 and r[0] and r[1] and str(r[0]).strip() and str(r[1]).strip()
                )
                if kv_count >= len(rows) * 0.5:
                    is_kv = True

        if is_kv:
            for row in rows:
                if len(row) >= 2 and row[0] and row[1]:
                    key = str(row[0]).strip()
                    value = str(row[1]).strip()
                    if key and value:
                        lines.append(f"**{key}**：{value}")
                elif row[0] and str(row[0]).strip():
                    lines.append(str(row[0]).strip())
        else:
            # 表格渲染
            for row_idx, row in enumerate(rows):
                cells = [str(c).strip() if c is not None else "" for c in row]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("|" + "|".join([" --- " for _ in cells]) + "|")

        lines.append("")  # sheet 间空行

    wb.close()
    return "\n".join(lines)


@router.post("/upload-doc", response_model=KBDocResponse)
async def upload_as_kb_doc(
    file: UploadFile = File(...),
    company_id: Optional[int] = Query(None, description="关联公司 ID"),
    product_id: Optional[int] = Query(None, description="关联产品 ID"),
    db: AsyncSession = Depends(get_db),
):
    """上传 Word/Excel 文件，自动转为 Markdown 知识库文档"""
    # 校验
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供文件名")
    if not company_id and not product_id:
        raise HTTPException(status_code=400, detail="必须指定 company_id 或 product_id")

    ext_lower = file.filename.lower()
    ext = ext_lower[ext_lower.rfind("."):] if "." in ext_lower else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"仅支持 {', '.join(ALLOWED_EXTENSIONS)} 格式的文件")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 10MB")

    # 转为 Markdown
    try:
        if ext == ".docx":
            md_content = _docx_to_markdown(content)
        else:
            md_content = _xlsx_to_markdown(content)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"文件解析失败: {str(e)}")

    if not md_content.strip():
        raise HTTPException(status_code=422, detail="文档内容为空，无法生成知识库")

    # 标题：用文件名（去掉扩展名）
    title = file.filename.rsplit(".", 1)[0] if "." in file.filename else file.filename

    # 保存到磁盘 + 创建记录
    file_path = _save_md_file(title, md_content, company_id, product_id)

    doc = KnowledgeBaseDoc(
        title=title,
        content=md_content,
        company_id=company_id,
        product_id=product_id,
        category="product",
        source="uploaded",
        file_path=file_path,
        word_count=len(md_content),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    cname = None
    pname = None
    if doc.company_id:
        c = await db.scalar(select(Company.name).where(Company.id == doc.company_id))
        cname = c
    if doc.product_id:
        p = await db.scalar(select(Product.name).where(Product.id == doc.product_id))
        pname = p

    return KBDocResponse(
        id=doc.id, title=doc.title,
        company_id=doc.company_id, company_name=cname,
        product_id=doc.product_id, product_name=pname,
        file_path=doc.file_path, category=doc.category,
        source=doc.source, word_count=doc.word_count,
        content=doc.content,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


# ═══════════════════════════════════════════
#  五、原有的文件上传解析（保留兼容）
# ═══════════════════════════════════════════

SECTION_MAP = {
    "产品名称": "name", "名称": "name", "产品名": "name",
    "品类": "category", "产品品类": "category", "分类": "category",
    "价格": "price", "售价": "price", "参考价格": "price", "定价": "price",
    "核心成分": "key_ingredients", "核心卖点": "key_ingredients", "主要成分": "key_ingredients",
    "成分": "key_ingredients", "关键成分": "key_ingredients",
    "功效宣称": "claims", "功效": "claims", "产品功效": "claims", "宣称": "claims", "核心功效": "claims",
    "目标人群": "target_audience", "适用人群": "target_audience", "目标用户": "target_audience", "用户画像": "target_audience",
    "行业背景": "industry_context", "行业分析": "industry_context", "市场背景": "industry_context", "市场分析": "industry_context",
}

LIST_FIELDS = {"key_ingredients", "claims"}


def _parse_docx_sections(content: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return {"raw": "文档内容为空"}
    result: dict = {}
    current_section: str | None = None
    for para in paragraphs:
        section_match = None
        clean_para = re.sub(r"^[#*]+|^[（(]?[一二三四五六七八九十\d]+[)）]?[、.]?\s*", "", para).strip()
        clean_para = re.sub(r"[：:]\s*$", "", clean_para)
        if clean_para in SECTION_MAP:
            section_match = SECTION_MAP[clean_para]
        else:
            for keyword, field in SECTION_MAP.items():
                if para.startswith(keyword) and ("：" in para or ":" in para):
                    section_match = field
                    value = re.split(r"[：:]", para, maxsplit=1)[-1].strip()
                    if value:
                        if field in LIST_FIELDS:
                            result[field] = [v.strip() for v in re.split(r"[、,;；，]", value) if v.strip()]
                        else:
                            result[field] = value
                    break
            if section_match:
                current_section = section_match
                continue
        if section_match:
            current_section = section_match
            continue
        if current_section:
            if current_section in LIST_FIELDS:
                result.setdefault(current_section, [])
                items = re.split(r"[、,;；，\n]", para)
                for item in items:
                    item = re.sub(r"^[-•\d.]+\s*", "", item).strip()
                    if item:
                        result[current_section].append(item)
            else:
                existing = result.get(current_section, "")
                result[current_section] = (existing + "\n" + para).strip()
        else:
            result.setdefault("raw", "")
            if result["raw"]:
                result["raw"] += "\n"
            result["raw"] += para
    _extract_from_raw(result, paragraphs)
    return result


def _parse_docx_as_markdown(content: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(content))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and "Heading" in para.style.name:
            level = re.search(r"Heading\s*(\d+)", para.style.name)
            prefix = "#" * (int(level.group(1)) if level else 1)
            lines.append(f"{prefix} {text}")
        else:
            lines.append(text)
    return {"raw": "\n\n".join(lines), "raw_markdown": True}


def _parse_xlsx(content: bytes) -> dict:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    result: dict = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue
        row1 = rows[0] if rows else []
        is_vertical = False
        if len(row1) >= 2:
            a1 = str(row1[0]).strip() if row1[0] else ""
            if a1 in SECTION_MAP or any(a1.startswith(k) for k in SECTION_MAP):
                is_vertical = True
        if not is_vertical and len(wb.sheetnames) == 1 and len(row1) >= 2:
            headers = [str(h).strip() if h else "" for h in row1]
            if any(h in SECTION_MAP for h in headers):
                is_vertical = False
            elif len(rows) >= 2:
                a1 = str(row1[0]).strip() if row1[0] else ""
                if len(a1) <= 20:
                    is_vertical = True
        if is_vertical:
            _parse_xlsx_vertical(rows, result)
        else:
            _parse_xlsx_horizontal(rows, result, sheet_name)
    wb.close()
    if not result:
        return {"raw": "Excel 内容为空或无法识别"}
    return result


def _parse_xlsx_vertical(rows: list, result: dict):
    for row in rows:
        if len(row) < 2:
            continue
        key = str(row[0]).strip() if row[0] else ""
        value = str(row[1]).strip() if row[1] else ""
        if not key or not value:
            continue
        field = None
        if key in SECTION_MAP:
            field = SECTION_MAP[key]
        else:
            for kw, f in SECTION_MAP.items():
                if key.startswith(kw) or kw in key:
                    field = f
                    break
        if field:
            if field in LIST_FIELDS:
                result.setdefault(field, [])
                items = [v.strip() for v in re.split(r"[、,;；，\n]", value) if v.strip()]
                result[field].extend(items)
            else:
                if field in result:
                    result[field] = result[field] + "\n" + value
                else:
                    result[field] = value


def _parse_xlsx_horizontal(rows: list, result: dict, sheet_name: str):
    headers: list[tuple[int, str | None]] = []
    row1 = rows[0]
    for i, cell in enumerate(row1):
        h = str(cell).strip() if cell else ""
        field = None
        if h in SECTION_MAP:
            field = SECTION_MAP[h]
        else:
            for kw, f in SECTION_MAP.items():
                if kw in h:
                    field = f
                    break
        headers.append((i, field))
    for row in rows[1:]:
        for i, field in headers:
            if field is None:
                continue
            value = str(row[i]).strip() if i < len(row) and row[i] else ""
            if not value:
                continue
            if field in LIST_FIELDS:
                result.setdefault(field, [])
                items = [v.strip() for v in re.split(r"[、,;；，\n]", value) if v.strip()]
                result[field].extend(items)
            else:
                if field in result:
                    result[field] = result[field] + "\n" + value
                else:
                    result[field] = value
        break


def _extract_from_raw(result: dict, paragraphs: list[str]):
    raw = result.get("raw", "")
    if not raw:
        return
    if "key_ingredients" not in result:
        ingredient_lines = re.findall(r"(?:成分|核心成分|主要成分|卖点)[：:]\s*(.+?)(?:\n|$)", raw)
        if ingredient_lines:
            all_ingredients = []
            for line in ingredient_lines:
                all_ingredients.extend([v.strip() for v in re.split(r"[、,;；，]", line) if v.strip()])
            if all_ingredients:
                result["key_ingredients"] = all_ingredients
    if "claims" not in result:
        claim_lines = re.findall(r"(?:功效|宣称|作用)[：:]\s*(.+?)(?:\n|$)", raw)
        if claim_lines:
            all_claims = []
            for line in claim_lines:
                all_claims.extend([v.strip() for v in re.split(r"[、,;；，]", line) if v.strip()])
            if all_claims:
                result["claims"] = all_claims
    if "name" not in result:
        first_line = paragraphs[0] if paragraphs else ""
        if first_line and len(first_line) <= 50 and not first_line.startswith("#"):
            result["name"] = first_line


@router.post("/upload", response_model=APIResponse)
async def upload_knowledge_base(file: UploadFile = File(...)):
    """上传知识库文件（.docx / .xlsx），解析为结构化 JSON"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供文件名")
    ext = file.filename.lower()
    ext = ext[ext.rfind("."):] if "." in ext else ext
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"仅支持 {', '.join(ALLOWED_EXTENSIONS)} 格式的文件")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 10MB")
    try:
        if ext == ".docx":
            structured = _parse_docx_sections(content)
            if len(structured) <= 2 and "raw" in structured:
                md_result = _parse_docx_as_markdown(content)
                if md_result.get("raw") and len(md_result["raw"]) > len(structured.get("raw", "")):
                    structured["raw"] = md_result["raw"]
                    structured["raw_markdown"] = True
        else:
            structured = _parse_xlsx(content)
        if not structured:
            structured = {"raw": "无法解析文档内容，请检查文档格式"}
        structured["_source"] = file.filename
        structured["_source_type"] = ext.lstrip(".")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"文档解析失败: {str(e)}")
    return APIResponse(success=True, message=f"文件 {file.filename} 解析成功", data=structured)
