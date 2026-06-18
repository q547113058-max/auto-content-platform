from docx import Document

doc = Document()
doc.add_heading('晶透焕亮精华液', level=1)
p = doc.add_paragraph()
p.add_run('品类：').bold = True
p.add_run('护肤品')
p = doc.add_paragraph()
p.add_run('参考价格：').bold = True
p.add_run('289元/30ml')
doc.add_heading('核心成分', level=2)
doc.add_paragraph('烟酰胺 5%')
doc.add_paragraph('透明质酸钠')
doc.add_paragraph('熊果苷提取物')
doc.add_paragraph('维生素C衍生物')
doc.add_paragraph('积雪草提取物')
doc.add_heading('功效宣称', level=2)
doc.add_paragraph('提亮肤色、淡化暗沉')
doc.add_paragraph('补水保湿、修护屏障')
doc.add_paragraph('改善粗糙、细腻毛孔')
p = doc.add_paragraph()
p.add_run('目标人群：').bold = True
p.add_run('25-40岁都市女性，肤色暗沉、熬夜党、初抗老需求')
doc.add_heading('行业背景', level=2)
doc.add_paragraph('2024年中国功效护肤品市场规模突破1200亿元，美白淡斑品类年增长超25%。')
doc.save('test_knowledge.docx')
print('OK')
